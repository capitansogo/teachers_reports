import locale

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')

data = [
    ['02.12', 'Сценарные языки программирования (ИСиТ)', 'ЦПИ-31', 8, 'Лаб.'],
    ['04.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['05.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['17.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['09.12', 'Сценарные языки программирования (ИСиТ)', 'ЦПИ-31', 8, 'Лаб.'],
    ['13.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['14.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['16.12', 'Сценарные языки программирования (ИСиТ)', 'ЦПИ-31', 8, 'Лаб.'],
    ['21.12', 'Работа с дипломниками', 'ДСИТ-47', 4],
    ['23.12', 'Сценарные языки программирования (ИСиТ)', 'ЦПИ-31', 8, 'Лаб.'],

]


def lengs(string):
    one = '_' * ((40 - len(string)) // 2)
    two = '_' * ((40 - len(string)) // 2)
    return one, two


def create_report(fio, dolzhnost, podrazdelenie, data, date_start, date_end, lessons):
    document = Document()

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    table = document.add_table(rows=1, cols=2)
    paragraph_format = table.cell(0, 0).paragraphs[0].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    cell = table.rows[0].cells[1]

    cell.text = "Проректору по ОД и ВР"
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    p2 = cell.add_paragraph("Голкиной В.А.")
    p2.style.font.name = "Times New Roman"
    p2.style.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(0)
    p2.style.font.name = "Times New Roman"
    p2.style.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(0)
    p3 = cell.add_paragraph()
    one, two = lengs(dolzhnost)
    run = p3.add_run(f'{one}')
    run = p3.add_run(f'{dolzhnost}')
    run.underline = True
    run = p3.add_run(f'{two}')
    p3.paragraph_format.space_after = Pt(0)
    p3.style.font.name = "Times New Roman"
    p3.style.font.size = Pt(12)
    p4 = cell.add_paragraph()
    run = p4.add_run("(должность преподавателя)")
    run.font.superscript = True
    p4.paragraph_format.alignment = 1
    p4.paragraph_format.space_after = Pt(0)
    p6 = cell.add_paragraph()
    one, two = lengs(podrazdelenie)
    run = p6.add_run(f'{one}')
    run = p6.add_run(f'{podrazdelenie}')
    run.underline = True
    run = p6.add_run(f'{two}')
    p6.paragraph_format.space_after = Pt(0)
    p6.style.font.name = "Times New Roman"
    p6.style.font.size = Pt(12)
    p7 = cell.add_paragraph()
    run = p7.add_run('(структурное подразделение, если совместитель, то указать)')
    run.font.superscript = True
    p7.paragraph_format.space_after = Pt(0)
    p7.paragraph_format.alignment = 1
    p7.style.font.name = "Times New Roman"
    p7.style.font.size = Pt(12)
    p9 = cell.add_paragraph()
    one, two = lengs(fio)
    run = p9.add_run(f'{one}')
    run = p9.add_run(f'{fio}')
    run.underline = True
    run = p9.add_run(f'{two}')
    p9.paragraph_format.space_after = Pt(0)
    p9.style.font.name = "Times New Roman"
    p9.style.font.size = Pt(12)
    p10 = cell.add_paragraph()
    run = p10.add_run('(ФИО сотрудника)')
    run.font.superscript = True
    p10.paragraph_format.alignment = 1
    p10.paragraph_format.space_after = Pt(0)
    p10.style.font.name = "Times New Roman"
    p10.style.font.size = Pt(12)

    p = document.add_paragraph("Заявление")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    run = document.add_paragraph("Мною в течение с ")
    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run.paragraph_format.space_after = Pt(0)
    run.paragraph_format.first_line_indent = Pt(36)  # Размер красной строки
    run.add_run("«")
    run.add_run(f'{date_start.strftime("%d")}').underline = True
    run.add_run("» ")
    run.add_run(f'{(date_start.strftime("%B")[0:-1]).lower()}я').underline = True
    run.add_run(" по ")
    run.add_run("«")
    run.add_run(f'{date_end.strftime("%d")}').underline = True
    run.add_run("» ")
    run.add_run(f'{(date_end.strftime("%B")[0:-1]).lower()}я {date_end.strftime("%Y")} г.').underline = True
    run.add_run(" проведены на кафедре ")
    run.add_run("Информационные системы и технологии").underline = True
    run.add_run(" занятия по дисциплинам ")
    for i in range(len(lessons)):
        run.add_run(f'{lessons[i]}').underline = True
        if i != len(lessons) - 1:
            run.add_run(", ")
    run.add_run(" со студентами следующие занятия:")
    run.add_run()

    table = document.add_table(rows=len(data) + 2, cols=5)
    table.style = 'Table Grid'
    table.autofit = True
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).text = "Дата"
    table.cell(0, 1).text = "Наименование учебных занятий"
    table.cell(0, 2).text = "Группа"
    table.cell(0, 3).text = "Кол-во академ. часов"
    table.cell(0, 4).text = "Примечание"
    table.cell(0, 0).width = Cm(1.2)
    table.cell(0, 1).width = Cm(6)
    table.cell(0, 2).width = Cm(1.2)
    table.cell(0, 3).width = Cm(1.2)
    table.cell(0, 4).width = Cm(1.2)

    header_row = table.rows[0]
    for cell in header_row.cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(11)

    for i in range(len(data)):
        for j in range(len(data[i])):
            cell = table.cell(i + 1, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell.text = str(data[i][j])
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)  # Размер шрифта по умолчанию

            # Выравнивание по левому краю для 2-й колонки
            if j == 1:
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Объединение первого и второго столбцов в последней строке
    merged_cell = table.cell(len(data) + 1, 0).merge(table.cell(len(data) + 1, 1).merge(table.cell(len(data) + 1, 2)))
    merged_cell.text = "ВСЕГО:"
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    last_row = table.rows[-1]

    for cell in last_row.cells:
        cell.borders = None

    print(data)
    sum = 0
    for i in range(len(data)):
        if data[i][3] is not None:
            sum += data[i][3]

    table.cell(len(data) + 1, 3).text = str(sum)
    cell = table.cell(len(data) + 1, 3)
    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph()

    # Добавление таблицы
    table = document.add_table(rows=2, cols=2)

    # Заполнение первой строки
    cell_1_1 = table.cell(0, 0)
    cell_1_1.text = "Занятия проведены"
    cell_1_1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell_1_2 = table.cell(0, 1)
    cell_1_2_paragraph = cell_1_2.paragraphs[0]
    cell_1_2_paragraph.add_run("_" * 40).bold = True

    cell_1_2_paragraph.add_run("\n(подпись преподавателя)")
    cell_1_2_paragraph.runs[1].font.superscript = True
    cell_1_2_paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
    cell_1_2_paragraph.paragraph_format.space_after = Pt(0)

    # Заполнение первой строки
    cell_1_1 = table.cell(1, 0)
    cell_1_1.text = "Проведение занятий подтверждаю"
    cell_1_1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell_1_2 = table.cell(1, 1)
    cell_1_2_paragraph = cell_1_2.paragraphs[0]
    cell_1_2_paragraph.add_run("_" * 40).bold = True

    cell_1_2_paragraph.add_run("\n(подпись зав. кафедрой)")
    cell_1_2_paragraph.runs[1].font.superscript = True
    cell_1_2_paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
    cell_1_2_paragraph.paragraph_format.space_after = Pt(0)

    p11 = document.add_paragraph()
    run = p11.add_run('«____»_________________________________20_____г.')
    p11.paragraph_format.space_after = Pt(0)

    p12 = document.add_paragraph()
    run = p12.add_run('Заключение планово-финансового управления')
    run.font.bold = True
    run.font.size = Pt(12)
    p12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавление таблицы 3 строки 2 столбца
    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0, 0).text = "Количество часов"
    table.cell(1, 0).text = "Оплата в час"
    table.cell(2, 0).text = "ИТОГО, руб."

    p13 = document.add_paragraph('Оплата за счет почасового фонда (рублей)')
    p13.paragraph_format.space_after = Pt(0)

    document.add_paragraph("_" * 72)

    p14 = document.add_paragraph('ПФУ')
    run = p14.add_run('_' * 30 + ' ' + '_' * 14 + ' ' + '_' * 18)
    p14.paragraph_format.space_after = Pt(0)

    p15 = document.add_paragraph()
    run = p15.add_run(' ' * 30)
    run = p15.add_run('(вид средств)')
    run.font.superscript = True
    run = p15.add_run(' ' * 35)
    run = p15.add_run('(подпись)')
    run.font.superscript = True
    run = p15.add_run(' ' * 14)
    run = p15.add_run('(расшифровка подписи)')
    run.font.superscript = True

    document.save('demo.docx')

    return document
